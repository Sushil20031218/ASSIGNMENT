import json
from docx import Document
from docx.shared import Inches
from openpyxl.workbook import Workbook
#all the libraries used to create docx file, exel file, imprt json and for the size of image

def main():#creating main function
    try:
        with open('curstomer_dataset.json', 'r') as file: #accessing the provided json file in read mode
            transactions = json.load(file) #reading the json file

    except FileNotFoundError:
        print("file not found.")
        return #this block is for checking whether the code block is running or not
    except json.JSONDecodeError:
        print("There was an error decoding the JSON file. Check its structure.")
        return

    for transaction in transactions: # access the invoice  if the invoice no. contanis H1 in it
        if "H1" in transaction['invoice_number']:
            doc = Document()
            try:
                doc.add_picture('logo.png', width=Inches(2))
                print('image') #add image as logo
            except Exception as e:
                print(f"An error occurred while adding the logo: {e}")
#inserting all the required credientials
            doc.add_heading('Invoice', level=1)
            doc.add_paragraph(f"Customer name: {transaction['customer_name']}")
            doc.add_paragraph(f"Customer phone: {transaction['customer_phone']}")
            doc.add_paragraph(f"Customer address: {transaction['customer_address']}")
            doc.add_paragraph(f"Invoice number: {transaction['invoice_number']}")
            doc.add_paragraph(f"Invoice date: {transaction['invoice_date']}")
#accessing the items dictonary from the json
            items = transaction.get('items_dict', {})
            table = doc.add_table(rows=1, cols=3) #created the table for
            table.style = 'Table Grid' #determine the style of table
            hdr_cells = table.rows[0].cells
#determinig the place for serial no, item name and price of the product
            for index, (item, price) in enumerate(items.items(), start=1): #
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                row_cells[1].text = item
                row_cells[2].text = str(price)


            filename = f'invoice_{transaction["invoice_number"]}.docx'

            num_items = 0
            subtotal = 0
            tax_rate = 0.13  # 13%
            tax_amount = 0
            #initializing the initial value

            # Calculate the number of items, subtotal, and tax
            for row in table.rows[1:]:  # Skip the header row
                num_items += 1
                price = float(row.cells[2].text)
                subtotal += price

            #calculating the tax amt
            tax_amount = subtotal * tax_rate

            # Calculate total (rounded)
            total = round(subtotal + tax_amount, 2)

            # Print the calculated values
            doc.add_paragraph(f"Number of Items: {num_items}")
            doc.add_paragraph(f"Subtotal: {subtotal}")
            doc.add_paragraph(f"Tax (13%): {tax_amount}")
            doc.add_paragraph(f"Total (Rounded): {total}")


            filename = f'invoice_{transaction["invoice_number"]}.docx' #generating random file name

            doc.save(filename)
            print(f"Invoice saved as {filename}.")

            wb = Workbook()
            ws = wb.active




#acessing the data from the transactions which contains H1 in its invoice and mentioning yes if it contains H1 in invoice id
            for transaction in transactions:
                contains_h1 = 'Yes' if 'H1' in transaction['invoice_number'] else ''
                data_row = [
                    transaction.get('customer_name', ''),
                    transaction.get('customer_phone', ''),
                    transaction.get('customer_address', ''),
                    transaction.get('invoice_number', ''),
                    transaction.get('invoice_date', ''),
                    contains_h1
                ]
                ws.append(data_row)

#saving the exel file
            wb.save('invoices_report.xlsx')
            print("Excel report generated.")
if __name__ == '__main__':
    main() #calling main function